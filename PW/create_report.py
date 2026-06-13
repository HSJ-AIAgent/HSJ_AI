from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 페이지 여백 설정 ──
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ── 헬퍼: 단락 아래 테두리 ──
def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── 헬퍼: 음영 배경 ──
def set_shading(paragraph, fill_color='EBF3FB'):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)

# ════════════════════════════════════════
#  표지
# ════════════════════════════════════════
# 상단 여백
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('인공지능(AI)')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('위키백과 요약 보고서')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

# 구분선
line_p = doc.add_paragraph()
line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = line_p.add_run('─' * 40)
run3.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

# 메타 정보 박스
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_shading(meta_p, 'EBF3FB')
meta_run = meta_p.add_run(
    f'출처: 위키백과 한국어판  |  URL: ko.wikipedia.org/wiki/인공지능\n'
    f'작성일: {datetime.date.today().strftime("%Y년 %m월 %d일")}  |  언어 지원: 177개 언어'
)
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

for _ in range(6):
    doc.add_paragraph()

# ════════════════════════════════════════
#  1. 개요
# ════════════════════════════════════════
doc.add_page_break()

h1 = doc.add_paragraph()
add_bottom_border(h1)
r = h1.add_run('1. 개요')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

overview_text = (
    '인공지능(人工智能, Artificial Intelligence, AI)은 인간의 학습능력, 추론능력, 지각능력을 '
    '인공적으로 구현하려는 컴퓨터 과학의 세부 분야이다. 정보공학 분야의 핵심 인프라 기술로 자리잡고 있으며, '
    '인간을 포함한 동물이 갖고 있는 자연 지능(natural intelligence)과는 구별되는 개념이다.\n\n'
    '인간의 지능을 모방한 기능을 갖춘 컴퓨터 시스템으로, 인간의 지능을 기계에 인공적으로 구현한 것이다. '
    '이 용어는 지능을 만들 수 있는 방법론이나 실현 가능성을 연구하는 과학 기술 분야를 지칭하기도 한다.\n\n'
    '초기 인공지능의 대표적인 정의는 다트머스 회의에서 존 매카시(John McCarthy)가 제안한 '
    '"기계를 인간 행동의 지식에서와 같이 행동하게 만드는 것"이다. 대부분의 정의는 '
    '① 인간처럼 사고하는 시스템, ② 인간처럼 행동하는 시스템, '
    '③ 이성적으로 사고하는 시스템, ④ 이성적으로 행동하는 시스템의 4가지로 분류된다.'
)
p = doc.add_paragraph(overview_text)
p.paragraph_format.first_line_indent = Cm(0.5)
p.runs[0].font.size = Pt(10.5)

# ════════════════════════════════════════
#  2. 인공지능의 분류
# ════════════════════════════════════════
doc.add_paragraph()
h2 = doc.add_paragraph()
add_bottom_border(h2)
r = h2.add_run('2. 인공지능의 분류')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

# 약인공지능
weak_title = doc.add_paragraph()
rw = weak_title.add_run('▶ 약인공지능 (Weak AI)')
rw.bold = True
rw.font.size = Pt(11)
rw.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

weak_p = doc.add_paragraph(
    '사진에서 물체를 찾거나 소리를 듣고 상황을 파악하는 등 특정 문제 해결에 특화된 인공지능이다. '
    '현실적이고 실용적인 목표를 갖고 개발되며, 일반적인 지능이 아닌 특정 도구로 활용된다. '
    '오늘날 상용화된 AI 서비스(음성인식, 이미지 분류, 추천 시스템 등)의 대부분이 이에 해당한다.'
)
weak_p.paragraph_format.left_indent = Cm(0.7)
weak_p.runs[0].font.size = Pt(10.5)

doc.add_paragraph()

# 강인공지능
strong_title = doc.add_paragraph()
rs = strong_title.add_run('▶ 강인공지능 (Strong AI / AGI)')
rs.bold = True
rs.font.size = Pt(11)
rs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

strong_p = doc.add_paragraph(
    '인간처럼 실제로 사고하여 문제를 해결할 수 있는 "일반 지능(AGI, Artificial General Intelligence)"을 '
    '인공적으로 구현하려는 시도이다. AGI 구현을 위해서는 추론, 문제 해결, 지식 표현, 계획 수립, '
    '의사결정, 학습이라는 지능적 특성들이 통합적으로 작용해야 하며, '
    '새로운 환경에서도 스스로 학습하여 적응하는 일반화 능력이 핵심 요소이다.'
)
strong_p.paragraph_format.left_indent = Cm(0.7)
strong_p.runs[0].font.size = Pt(10.5)

# ════════════════════════════════════════
#  3. 역사
# ════════════════════════════════════════
doc.add_paragraph()
h3 = doc.add_paragraph()
add_bottom_border(h3)
r = h3.add_run('3. 인공지능의 역사')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

history_data = [
    ('1943–1950년대', '이론적 기반 형성',
     '월터 피츠·워런 매컬러의 인공 신경망 연구, 앨런 튜링의 "계산 기계와 지능" 논문 발표(1950), '
     '튜링 테스트 제안. 마빈 민스키가 최초의 신경망 기계 SNARC 구축(1951).'),
    ('1956년', '인공지능 학문 탄생',
     '다트머스 회의에서 존 매카시가 "인공지능(Artificial Intelligence)"이라는 용어를 공식 제안하며 학문 분야로 확립.'),
    ('1956–1974년', '황금기',
     '초기 AI 프로그램들의 성과로 낙관적 기대가 고조. 자연어 처리, 문제 해결 프로그램 등이 등장.'),
    ('1974–1980년', '첫 번째 AI 겨울',
     '연산 능력의 한계와 복잡한 현실 문제 해결 실패로 연구비 삭감 및 침체기 돌입.'),
    ('1980–1987년', 'AI 붐 (전문가 시스템)',
     '전문가 시스템(Expert System)의 상업적 성공으로 AI 르네상스. 규칙 기반 추론 시스템이 산업계에 도입됨.'),
    ('1987–1993년', '두 번째 AI 겨울',
     '전문가 시스템의 유지 비용 문제 및 범용성 한계 노출로 재차 침체. 하드웨어 시장 붕괴도 영향.'),
    ('1993년–현재', '현대 AI의 발전',
     '머신러닝, 딥러닝, 빅데이터의 융합으로 혁신적 발전. 알파고(2016), ChatGPT(2022) 등 범용 AI 등장으로 '
     '다시 전 세계적인 AI 열풍.'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# 헤더
hdr_cells = table.rows[0].cells
headers = ['시기', '구분', '주요 내용']
for i, hdr in enumerate(headers):
    hdr_cells[i].text = hdr
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497D')
    hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

for period, label, desc in history_data:
    row_cells = table.add_row().cells
    row_cells[0].text = period
    row_cells[1].text = label
    row_cells[2].text = desc
    for cell in row_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)

# 열 너비 조정
from docx.shared import Cm
widths = [Cm(2.8), Cm(3.2), Cm(10.0)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

# ════════════════════════════════════════
#  4. 한계 및 문제점
# ════════════════════════════════════════
doc.add_paragraph()
h4 = doc.add_paragraph()
add_bottom_border(h4)
r = h4.add_run('4. 인공지능의 한계 및 문제점')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

issues = [
    ('거짓정보 전달 및 가짜뉴스', 'AI가 생성하는 콘텐츠의 신뢰성 문제. 딥페이크, 허위 정보 생성 등 사회적 혼란 야기 가능성.'),
    ('인간의 통제력 약화', 'AI 의사결정 과정의 불투명성(블랙박스 문제)으로 인한 자율 시스템에 대한 인간 통제 어려움.'),
    ('일자리 감소', '반복적·규칙적 업무의 자동화로 특정 직군의 일자리 대체 우려. 노동 시장 구조적 변화 초래.'),
    ('윤리 문제', '알고리즘 편향(bias), 차별적 의사결정, AI의 도덕적 판단 기준 부재 등 윤리적 쟁점 내포.'),
    ('개인정보 유출', '대규모 데이터 학습 과정에서의 개인정보 침해, 프라이버시 위협 가능성.'),
    ('사고력 저하', 'AI 의존도 증가로 인한 인간의 비판적 사고력 및 창의적 문제 해결 능력 저하 우려.'),
]

for i, (issue, desc) in enumerate(issues):
    fill = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
    issue_p = doc.add_paragraph()
    set_shading(issue_p, fill)
    issue_p.paragraph_format.left_indent = Cm(0.3)
    ri = issue_p.add_run(f'● {issue}  ')
    ri.bold = True
    ri.font.size = Pt(10.5)
    ri.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    rd = issue_p.add_run(desc)
    rd.font.size = Pt(10)

# ════════════════════════════════════════
#  5. 주요 응용 분야
# ════════════════════════════════════════
doc.add_paragraph()
h5 = doc.add_paragraph()
add_bottom_border(h5)
r = h5.add_run('5. 주요 응용 분야')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

apps = [
    '자연어 처리 (NLP): 기계 번역, 챗봇, 음성인식, 텍스트 요약 등',
    '컴퓨터 비전: 이미지·영상 인식, 자율주행, 의료 영상 진단',
    '전문가 시스템: 의료 진단, 금융 분석, 법률 자문 지원',
    '로보틱스: 산업용 자동화 로봇, 드론, 물류 자동화',
    '추천 시스템: 콘텐츠 개인화, 전자상거래, 광고 타겟팅',
    '게임 AI: 체스·바둑 AI(알파고), 실시간 전략 게임 에이전트',
    '창작 AI: 음악·미술·글쓰기 생성 (ChatGPT, DALL-E 등)',
]

for app in apps:
    ap = doc.add_paragraph(style='List Bullet')
    ar = ap.add_run(app)
    ar.font.size = Pt(10.5)

# ════════════════════════════════════════
#  6. 미래 전망
# ════════════════════════════════════════
doc.add_paragraph()
h6 = doc.add_paragraph()
add_bottom_border(h6)
r = h6.add_run('6. 미래 전망')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

future_p = doc.add_paragraph(
    '인공지능의 미래는 초지능(Superintelligence)의 등장 가능성과 그에 따른 위험성 논의가 핵심이다. '
    '닉 보스트롬 등 미래학자들은 인간 지능을 초월하는 AGI가 실현될 경우, '
    '통제 불가능한 자율 의사결정 시스템으로 인류에게 실존적 위험을 초래할 수 있다고 경고한다.\n\n'
    '반면 낙관적 관점에서는 AI가 기후 변화, 질병 치료, 빈곤 해소 등 인류의 난제를 해결하는 '
    '핵심 도구가 될 것으로 기대한다. 현재는 AI 안전성(AI Safety), 정렬 문제(Alignment Problem), '
    '글로벌 AI 거버넌스 구축이 전 세계적 과제로 대두되고 있다.'
)
future_p.paragraph_format.first_line_indent = Cm(0.5)
future_p.runs[0].font.size = Pt(10.5)

# ════════════════════════════════════════
#  7. 유명 인물 및 자매 프로젝트
# ════════════════════════════════════════
doc.add_paragraph()
h7 = doc.add_paragraph()
add_bottom_border(h7)
r = h7.add_run('7. 인공지능 관련 주요 인물')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

persons = [
    ('앨런 튜링 (Alan Turing)', '튜링 테스트 제안, 현대 컴퓨터 과학의 아버지'),
    ('존 매카시 (John McCarthy)', '"인공지능" 용어 창시, 다트머스 회의 주도'),
    ('마빈 민스키 (Marvin Minsky)', '최초 신경망 기계 SNARC 개발, MIT AI 연구소 공동 창립'),
    ('제프리 힌튼 (Geoffrey Hinton)', '딥러닝의 아버지, 역전파 알고리즘 발전에 기여'),
    ('얀 르쿤 (Yann LeCun)', '합성곱 신경망(CNN) 개발, 컴퓨터 비전 혁신'),
    ('앤드류 응 (Andrew Ng)', '딥러닝 교육 대중화, Coursera AI 강좌 개설'),
]

for name, desc in persons:
    pp = doc.add_paragraph()
    pp.paragraph_format.left_indent = Cm(0.3)
    pn = pp.add_run(f'■ {name}: ')
    pn.bold = True
    pn.font.size = Pt(10.5)
    pn.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    pd = pp.add_run(desc)
    pd.font.size = Pt(10.5)

# ════════════════════════════════════════
#  푸터: 출처
# ════════════════════════════════════════
doc.add_paragraph()
footer_p = doc.add_paragraph()
set_shading(footer_p, 'D6E4F0')
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    '본 보고서는 위키백과 한국어판(ko.wikipedia.org)의 "인공지능" 문서를 기반으로 작성되었습니다.\n'
    f'작성일: {datetime.date.today().strftime("%Y년 %m월 %d일")}  |  '
    '라이선스: CC BY-SA 4.0'
)
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── 저장 ──
output_path = r'C:\Users\SBS\Desktop\HSJ_AI\PW\인공지능_위키백과_요약보고서.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
